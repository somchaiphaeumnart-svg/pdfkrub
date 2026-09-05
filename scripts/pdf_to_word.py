#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Advanced PDF to Word (.docx) Converter.
Supports two main engines:
1. Standard High-Fidelity Layout mode (using pdf2docx with LibreOffice fallback)
   Preserves Thai fonts, tables, shapes, colors, and layout structure.
2. Thai OCR mode (for scanned PDFs or image-based PDFs)
   Extracts text using PyMuPDF + Tesseract OCR (tha+eng) and outputs editable Word .docx.
"""

import sys
import os
import argparse
import subprocess
import tempfile
import zipfile
import xml.sax.saxutils as saxutils
import re
import unicodedata

def normalize_thai_text(text):
    if not text:
        return ""
    # PUA mapping
    pua_map = {
        0xF700: 0x0E10, 0xF701: 0x0E0D, 0xF702: 0x0E47, 0xF703: 0x0E34, 0xF704: 0x0E35,
        0xF705: 0x0E36, 0xF706: 0x0E37, 0xF70A: 0x0E48, 0xF70B: 0x0E49, 0xF70C: 0x0E4A,
        0xF70D: 0x0E4B, 0xF70E: 0x0E4C, 0xF710: 0x0E31, 0xF711: 0x0E34, 0xF712: 0x0E35,
        0xF713: 0x0E36, 0xF714: 0x0E37, 0xF715: 0x0E48, 0xF716: 0x0E49, 0xF717: 0x0E4A,
        0xF718: 0x0E4B, 0xF719: 0x0E4C, 0xF71A: 0x0E4D,
        0xF884: 0x0E47, 0xF885: 0x0E34, 0xF886: 0x0E35, 0xF887: 0x0E36, 0xF888: 0x0E37,
        0xF889: 0x0E48, 0xF88A: 0x0E49, 0xF88B: 0x0E4A, 0xF88C: 0x0E4B, 0xF88D: 0x0E4C,
        0xF894: 0x0E31, 0xF897: 0x0E48, 0xF898: 0x0E49, 0xF899: 0x0E4A, 0xF89A: 0x0E4B, 0xF89B: 0x0E4C
    }
    chars = [chr(pua_map.get(ord(c), ord(c))) for c in text]
    text = "".join(chars)

    # Mojibake detection & repair (CP874 misread as CP1252/Latin-1)
    def fix_mojibake(m):
        return "".join([chr(ord(c) - 0xA0 + 0x0E00) if 0xA1 <= ord(c) <= 0xFB else c for c in m.group(0)])
    text = re.sub(r'[\u00A1-\u00FB]{3,}', fix_mojibake, text)

    # 3-level Thai character stacking reordering
    text = re.sub(r'([\u0E48-\u0E4C])([\u0E34-\u0E37\u0E31\u0E47\u0E4D])', r'\2\1', text)
    text = re.sub(r'([\u0E48-\u0E4C])([\u0E38-\u0E3A])', r'\2\1', text)
    text = re.sub(r'\u0E4D\u0E32', '\u0E33', text)
    text = re.sub(r'([\u0E48-\u0E4B])\u0E4D\u0E32', r'\1\u0E33', text)
    text = re.sub(r'\u0E4D([\u0E48-\u0E4B])\u0E32', r'\1\u0E33', text)
    text = re.sub(r'\u0E33([\u0E48-\u0E4B])', r'\1\u0E33', text)
    text = re.sub(r'([\u0E31\u0E34-\u0E3A\u0E47-\u0E4E])\1+', r'\1', text)

    # Accidental space removal
    text = re.sub(r'([เแโใไ])\s+([ก-ฮ])', r'\1\2', text)
    text = re.sub(r'([ก-ฮ])\s+([\u0E30-\u0E3A\u0E47-\u0E4E])', r'\1\2', text)
    text = re.sub(r'([ก-ฮ])\s+([าะำๅ])', r'\1\2', text)
    text = re.sub(r'([์])\s+([ก-ฮ])', r'\1\2', text)
    text = re.sub(r'[\uFFFD\u0000-\u0008\u000B-\u001F]', '', text)
    try:
        return unicodedata.normalize('NFC', text)
    except Exception:
        return text

def parse_page_range(range_str, total_pages):
    """
    Parse page range string like 'all', '1-3,5', '2,4,6'
    Returns 0-indexed list of page numbers.
    """
    if not range_str or range_str.strip().lower() == 'all':
        return list(range(total_pages))

    pages = set()
    parts = [p.strip() for p in range_str.split(',') if p.strip()]
    for part in parts:
        if '-' in part:
            bounds = part.split('-')
            if len(bounds) == 2:
                try:
                    start = int(bounds[0])
                    end = int(bounds[1])
                    for p in range(start, end + 1):
                        if 1 <= p <= total_pages:
                            pages.add(p - 1)
                except ValueError:
                    pass
        else:
            try:
                p = int(part)
                if 1 <= p <= total_pages:
                    pages.add(p - 1)
            except ValueError:
                pass

    sorted_pages = sorted(list(pages))
    return sorted_pages if sorted_pages else list(range(total_pages))

def get_pdf_page_count(pdf_path):
    """Get total page count of a PDF using PyMuPDF or pdfinfo."""
    try:
        import fitz
        doc = fitz.open(pdf_path)
        count = len(doc)
        doc.close()
        return count
    except Exception:
        pass

    try:
        res = subprocess.run(['pdfinfo', pdf_path], capture_output=True, text=True, timeout=15)
        if res.returncode == 0:
            for line in res.stdout.splitlines():
                if line.startswith('Pages:'):
                    return int(line.split(':', 1)[1].strip())
    except Exception:
        pass

    return 1

def build_docx_builtin(pages_text, output_path):
    """
    Generate a valid OpenXML .docx file using standard Python built-in modules.
    Ensures zero dependency failure.
    """
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    temp_file = output_path + ".tmp"

    body_xml_parts = []
    for p_idx, text in enumerate(pages_text):
        if p_idx > 0:
            # Page break
            body_xml_parts.append('<w:p><w:r><w:br w:type="page"/></w:r></w:p>')

        paragraphs = text.split("\n")
        for para in paragraphs:
            para_clean = para.strip()
            if not para_clean:
                body_xml_parts.append('<w:p><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>')
                continue
            escaped = saxutils.escape(para_clean)
            p_xml = (
                '<w:p>'
                '<w:pPr><w:spacing w:after="120"/><w:rPr><w:rFonts w:ascii="TH Sarabun New" w:hAnsi="TH Sarabun New" w:cs="TH Sarabun New"/><w:sz w:val="32"/><w:szCs w:val="32"/></w:rPr></w:pPr>'
                f'<w:r><w:rPr><w:rFonts w:ascii="TH Sarabun New" w:hAnsi="TH Sarabun New" w:cs="TH Sarabun New"/><w:sz w:val="32"/><w:szCs w:val="32"/></w:rPr><w:t>{escaped}</w:t></w:r>'
                '</w:p>'
            )
            body_xml_parts.append(p_xml)

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">\n'
        '<w:body>\n'
        + '\n'.join(body_xml_parts) +
        '\n<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr>\n'
        '</w:body>\n'
        '</w:document>'
    )

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">\n'
        '  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>\n'
        '  <Default Extension="xml" ContentType="application/xml"/>\n'
        '  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>\n'
        '</Types>'
    )

    pkg_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">\n'
        '  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>\n'
        '</Relationships>'
    )

    with zipfile.ZipFile(temp_file, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr('[Content_Types].xml', content_types)
        zf.writestr('_rels/.rels', pkg_rels)
        zf.writestr('word/document.xml', document_xml)

    if os.path.exists(output_path):
        os.remove(output_path)
    os.rename(temp_file, output_path)
    return True

def convert_with_pdf2docx(input_pdf, output_docx, page_indices=None):
    """Convert using pdf2docx for layout and Thai font preservation."""
    try:
        from pdf2docx import Converter
        cv = Converter(input_pdf)
        if page_indices is not None and len(page_indices) > 0:
            cv.convert(output_docx, pages=page_indices)
        else:
            cv.convert(output_docx)
        cv.close()
        return os.path.exists(output_docx) and os.path.getsize(output_docx) > 0
    except Exception as e:
        sys.stderr.write(f"pdf2docx notice: {e}\n")
        return False

def convert_with_libreoffice(input_pdf, output_docx):
    """Fallback conversion using LibreOffice headless."""
    try:
        out_dir = os.path.dirname(os.path.abspath(output_docx))
        base_name = os.path.splitext(os.path.basename(input_pdf))[0]

        cmd = [
            'soffice',
            '--headless',
            '--norestore',
            '--infilter=writer_pdf_import',
            '--convert-to', 'docx',
            '--outdir', out_dir,
            input_pdf
        ]
        res = subprocess.run(cmd, capture_output=True, timeout=120)
        generated = os.path.join(out_dir, f"{base_name}.docx")
        if os.path.exists(generated) and os.path.getsize(generated) > 0:
            if generated != output_docx:
                if os.path.exists(output_docx):
                    os.remove(output_docx)
                os.rename(generated, output_docx)
            return True
    except Exception as e:
        sys.stderr.write(f"LibreOffice notice: {e}\n")
    return False

def convert_with_ocr(input_pdf, output_docx, page_indices=None, ocr_lang="tha+eng"):
    """
    Render PDF pages to high-res images, perform Tesseract OCR,
    and generate editable Word document.
    """
    try:
        import fitz
        doc = fitz.open(input_pdf)
        total = len(doc)
        if page_indices is None or len(page_indices) == 0:
            target_pages = list(range(total))
        else:
            target_pages = [p for p in page_indices if 0 <= p < total]

        pages_text = []

        with tempfile.TemporaryDirectory() as tmp_dir:
            for idx, p_num in enumerate(target_pages):
                page = doc[p_num]
                pix = page.get_pixmap(dpi=200)
                img_path = os.path.join(tmp_dir, f"page_{p_num}.png")
                pix.save(img_path)

                out_base = os.path.join(tmp_dir, f"ocr_{p_num}")
                cmd = ['tesseract', img_path, out_base, '-l', ocr_lang, '--psm', '3', 'txt']
                res = subprocess.run(cmd, capture_output=True, timeout=90)
                txt_file = f"{out_base}.txt"
                if os.path.exists(txt_file):
                    with open(txt_file, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read().strip()
                else:
                    text = page.get_text().strip()

                text = normalize_thai_text(text)
                pages_text.append(text)

        doc.close()

        # Try building docx via python-docx if installed
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.oxml.ns import qn

            word_doc = Document()
            # Set Margins to 1 inch
            for section in word_doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            for idx, page_content in enumerate(pages_text):
                if idx > 0:
                    word_doc.add_page_break()

                for line in page_content.splitlines():
                    clean_line = line.strip()
                    p = word_doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(4)
                    p.paragraph_format.line_spacing = 1.15
                    if clean_line:
                        run = p.add_run(clean_line)
                        run.font.name = 'TH Sarabun New'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'TH Sarabun New')
                        run._element.rPr.rFonts.set(qn('w:cs'), 'TH Sarabun New')
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(30, 30, 30)

            word_doc.save(output_docx)
            if os.path.exists(output_docx) and os.path.getsize(output_docx) > 0:
                return True
        except Exception as e:
            sys.stderr.write(f"python-docx notice: {e}, falling back to built-in OpenXML writer\n")

        # Fallback to builtin XML generator
        return build_docx_builtin(pages_text, output_docx)

    except Exception as e:
        sys.stderr.write(f"OCR conversion notice: {e}\n")
        return False

def main():
    parser = argparse.ArgumentParser(description="Convert PDF to Word (.docx)")
    parser.add_argument("input_pdf", help="Input PDF file path")
    parser.add_argument("output_docx", help="Output DOCX file path")
    parser.add_argument("--mode", default="standard", choices=["standard", "ocr"], help="Conversion mode")
    parser.add_argument("--pages", default="all", help="Pages to convert (e.g. 'all', '1-3,5')")
    parser.add_argument("--ocr-lang", default="tha+eng", help="OCR language")
    parser.add_argument("--detect-tables", default="1", help="Detect tables (1/0)")
    parser.add_argument("--keep-images", default="1", help="Keep embedded images (1/0)")

    args = parser.parse_args()

    input_pdf = os.path.abspath(args.input_pdf)
    output_docx = os.path.abspath(args.output_docx)

    if not os.path.exists(input_pdf) or os.path.getsize(input_pdf) == 0:
        sys.stderr.write(f"Error: Input file {input_pdf} does not exist or is empty\n")
        sys.exit(1)

    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    total_pages = get_pdf_page_count(input_pdf)
    page_indices = parse_page_range(args.pages, total_pages)

    success = False

    if args.mode == "ocr":
        success = convert_with_ocr(input_pdf, output_docx, page_indices, args.ocr_lang)

    if not success and args.mode == "standard":
        success = convert_with_pdf2docx(input_pdf, output_docx, page_indices)
        if not success:
            success = convert_with_libreoffice(input_pdf, output_docx)
        if not success:
            success = convert_with_ocr(input_pdf, output_docx, page_indices, args.ocr_lang)

    if success and os.path.exists(output_docx) and os.path.getsize(output_docx) > 0:
        sys.exit(0)
    else:
        sys.stderr.write("PDF to Word conversion failed with all available engines\n")
        sys.exit(1)

if __name__ == '__main__':
    main()
